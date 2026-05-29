"""Extract plain text from a motion: DOCX, PDF, or .txt.

DOCX uses python-docx (paragraphs + tables). PDF uses pdfplumber. Both deps are
already standard in this space; PDF text extraction works on text-based PDFs —
scanned/image PDFs would need OCR, which is out of scope for v1.
"""

from __future__ import annotations

from pathlib import Path


def extract_text(path: str | Path) -> str:
    """Return the plain text of a document, dispatched by file extension."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _from_docx(path)
    if suffix == ".pdf":
        return _from_pdf(path)
    if suffix in (".txt", ".text", ".md"):
        return path.read_text(encoding="utf-8", errors="replace")
    raise ValueError(
        f"Unsupported document type '{suffix}'. Use .docx, .pdf, or .txt."
    )


def _from_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _from_pdf(path: Path) -> str:
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ""
            parts.append(txt)
    return "\n".join(parts)
